"""Render a draft document's HTML to a .docx file.

The draft CV/CL HTML uses a small, known vocabulary (from engine/draft.py's
prompt): <h3> for the name, <div class='role-h'> for section/role headers,
<p>, <ul><li>, <strong>/<b>, <em>/<i>, <a>, <br>. We walk it and emit a clean
Word document — name centered, a centered contact line, ruled section headers,
bullet lists — with email and LinkedIn/URLs turned into real hyperlinks so they
stay clickable, mirroring the reference layout's contact block and sections.
"""
from __future__ import annotations

import io
import re

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# email | http(s) url | bare www./linkedin.com path
TOKEN_RE = re.compile(
    r"[\w.+-]+@[\w-]+\.[\w.-]+"
    r"|https?://[^\s)<>]+"
    r"|(?:www\.|linkedin\.com/)[^\s)<>]+",
    re.I,
)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
ACCENT = RGBColor(0x2D, 0x6C, 0xDF)       # section-header accent
NAME_INK = RGBColor(0x16, 0x20, 0x2E)     # near-black for the name


def _spacing(run, twentieths: int) -> None:
    """Apply character spacing (in twentieths of a point) to a run."""
    rPr = run._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(twentieths))
    rPr.append(sp)


def _normalize_url(href: str) -> str:
    href = href.strip()
    if href.startswith(("http://", "https://", "mailto:")):
        return href
    if "@" in href and "/" not in href:
        return "mailto:" + href
    return "https://" + href


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _add_run(paragraph, text: str, bold: bool, italic: bool) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic


def _emit_text(paragraph, text: str, bold: bool, italic: bool) -> None:
    """Add text, auto-linking any email / URL / LinkedIn tokens inside it."""
    idx = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > idx:
            _add_run(paragraph, text[idx:m.start()], bold, italic)
        tok = m.group(0).rstrip(".,;)")
        _add_hyperlink(paragraph, _normalize_url(tok), tok)
        idx = m.start() + len(tok)
    if idx < len(text):
        _add_run(paragraph, text[idx:], bold, italic)


def _render_children(paragraph, element, bold=False, italic=False) -> None:
    for child in element.children:
        if isinstance(child, NavigableString):
            _emit_text(paragraph, str(child), bold, italic)
            continue
        nm = child.name
        if nm in ("strong", "b"):
            _render_children(paragraph, child, True, italic)
        elif nm in ("em", "i"):
            _render_children(paragraph, child, bold, True)
        elif nm == "a":
            text = child.get_text()
            href = child.get("href") or text
            _add_hyperlink(paragraph, _normalize_url(href), text)
        elif nm == "br":
            paragraph.add_run().add_break()
        else:  # mark, span, anything else — descend
            _render_children(paragraph, child, bold, italic)


def _section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    _spacing(r, 30)  # letter-spacing ~1.5pt for the uppercase header
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2D6CDF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def html_to_docx(html: str) -> bytes:
    soup = BeautifulSoup(html or "", "html.parser")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    name_done = False
    expect_contact = False
    for node in soup.children:
        if isinstance(node, NavigableString):
            if node.strip():
                _emit_text(doc.add_paragraph(), str(node).strip(), False, False)
            continue
        nm = node.name
        classes = node.get("class") or []
        if nm == "hr" and "pagebreak" in classes:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            continue
        is_heading = nm in ("h1", "h2", "h3")
        if is_heading and not name_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(node.get_text(strip=True))
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = NAME_INK
            p.paragraph_format.space_after = Pt(2)
            name_done = True
            expect_contact = True
        elif is_heading or (nm == "div" and "role-h" in classes):
            _section_heading(doc, node.get_text(strip=True))
            expect_contact = False
        elif nm == "ul":
            for li in node.find_all("li", recursive=False):
                _render_children(doc.add_paragraph(style="List Bullet"), li)
            expect_contact = False
        elif nm in ("p", "div"):
            p = doc.add_paragraph()
            if expect_contact and node.get_text(strip=True):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                expect_contact = False
            _render_children(p, node)
        else:
            _render_children(doc.add_paragraph(), node)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
